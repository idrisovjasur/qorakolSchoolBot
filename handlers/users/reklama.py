import asyncio
from aiogram.types import Message, ContentType, ReplyKeyboardRemove, CallbackQuery, InputFile
from data.config import ADMINS
from aiogram.dispatcher import FSMContext
from docx import Document
import os
from keyboards.default.contact import true_false, bosh_menu
from keyboards.inline.url_english import qora, how_subject
from loader import db, dp, bot
from aiogram.types import ChatActions


@dp.message_handler(text='/reklama_text', user_id=ADMINS[0])
@dp.message_handler(text='/reklama_text', user_id=ADMINS[1])
async def reklama_text_def(message: Message, state: FSMContext):
    await message.answer("reklama matnini kiriting!")
    await state.set_state('reklama_text')


@dp.message_handler(state='reklama_text')
async def reklama_def_dtat(message: Message, state: FSMContext):
    text = message.text
    users = db.select_all_users()
    await state.finish()
    for user in users:
        user_id = user[0]
        await bot.send_message(chat_id=user_id, text=text)
        await asyncio.sleep(0.05)
    await message.answer("Reklama hammaga yuborildi")


####RASM
@dp.message_handler(text='/reklama_photo', user_id=ADMINS[0])
@dp.message_handler(text='/reklama_photo', user_id=ADMINS[1])
async def reklama_photo_def(message: Message, state: FSMContext):
    await message.answer('Reklama Photo rasmini yuboring')
    await state.set_state('reklama_photo')


@dp.message_handler(state='reklama_photo', content_types=ContentType.PHOTO)
async def state_photo_defaa_data(message: Message, state: FSMContext):
    global photo_id
    photo_id = message.photo[-1].file_id
    await message.answer('Caption yuboring!')
    await state.set_state('caption')


@dp.message_handler(state='caption')
async def caption_def(message: Message, state: FSMContext):
    caption = message.text
    users = db.select_all_users()
    await state.finish()
    for user in users:
        user_id = user[0]
        await bot.send_photo(chat_id=user_id, photo=photo_id, caption=caption)
        await asyncio.sleep(0.05)

    await message.answer('reklama yuborildi hammaga')


@dp.message_handler(text='/reklama_vedio', user_id=ADMINS[0])
@dp.message_handler(text='/reklama_vedio', user_id=ADMINS[1])
async def video_def_statelar_bilan(message: Message, state: FSMContext):
    await message.answer('Reklama Video ni yuboring')
    await state.set_state('reklama_video')


@dp.message_handler(state='reklama_video', content_types=ContentType.VIDEO)
async def video_reklama_def(message: Message, state: FSMContext):
    global video_id
    video_id = message.video.file_id
    await message.answer('Caption yuboring')
    await state.set_state('caption_video')


@dp.message_handler(state='caption_video')
async def caption_video(message: Message, state: FSMContext):
    caption = message.text
    users = db.select_all_users()
    await state.finish()
    for user in users:
        user_id = user[0]
        await bot.send_video(chat_id=user_id, video=video_id, caption=caption)
        await asyncio.sleep(0.05)

    await message.answer('reklama hammaga yuborildi')


@dp.message_handler(text='/delete_users', user_id=ADMINS[1])
@dp.message_handler(text='/delete_users', user_id=ADMINS[0])
async def delete_user_def(message: Message, state: FSMContext):
    await message.answer('DIQQAT!\n'
                         'Hamma foydalanuvchi bazadan o\'chadi!', reply_markup=true_false)
    await state.set_state('delete_user')


@dp.callback_query_handler(state='delete_user')
async def delete_user_state_def(call: CallbackQuery, state: FSMContext):
    await call.message.delete_reply_markup()
    if call.data == '✅ Ha':
        db.delete_users()
        await call.answer("Hamma Foydalanuvchi bazadan o'chirildi ✅", show_alert=True)
        await call.message.answer("Bosh menu", reply_markup=bosh_menu)
        await state.finish()
    else:
        await call.answer('Bazada foydalanuvchilar o\'chirilmadi', show_alert=True)
        await state.finish()


@dp.callback_query_handler(state='how_subject')
async def add_techer_bazu(call: CallbackQuery, state: FSMContext):
    if call.data == 'ingliz':
        fan = 'english'
    elif call.data == 'math':
        fan = 'math'
    elif call.data == 'it':
        fan = 'it'
    await state.update_data(
        {
            'fan': fan
        }
    )
    await call.message.answer('Ism:')
    await state.set_state('teach_ism')


@dp.message_handler(state='teach_ism')
async def teach_ism_def(message: Message, state: FSMContext):
    ism = message.text
    await state.update_data(
        {
            "ism": ism
        }
    )
    await message.answer('Familya:')
    await state.set_state('teach_familya')


@dp.message_handler(state='teach_familya')
async def teach_fam_def(message: Message, state: FSMContext):
    familya = message.text
    await state.update_data(
        {
            'familya': familya
        }
    )
    await message.answer('Tel:')
    await state.set_state('teach_phone')


@dp.message_handler(state='teach_phone')
async def teach_phone_def(message: Message, state: FSMContext):
    phone = message.text
    await state.update_data(
        {
            'phone': phone
        }

    )
    await message.answer('Ielts:')
    await state.set_state('teach_ielts')


@dp.message_handler(state='teach_ielts')
async def teach_phone_def(message: Message, state: FSMContext):
    ielts = message.text
    await state.update_data(
        {
            'ielts': ielts
        }

    )
    await message.answer('Stage:')
    await state.set_state('teach_stage')


@dp.message_handler(state='teach_stage')
async def teach_phone_def(message: Message, state: FSMContext):
    stage = message.text
    await state.update_data(
        {
            'stage': stage
        }

    )
    await message.answer('Yosh:')
    await state.set_state('teach_yosh')


@dp.message_handler(state='teach_yosh')
async def teach_phone_def(message: Message, state: FSMContext):
    yosh = message.text
    await state.update_data(
        {
            'yosh': yosh
        }

    )
    data = await state.get_data()
    fan = data.get('fan')
    ism = data.get('ism')
    fam = data.get('familya')
    phone = data.get('phone')
    ielts = data.get('ielts')
    stage = data.get('stage')
    yosh = data.get('yosh')
    try:
        db.add_teacher(
            science=fan,
            first_name=ism,
            last_name=fam,
            phone=phone,
            ielts=ielts,
            stage=stage,
            age=yosh
        )
        await message.answer('Qabul qilindi!')
    except:
        await message.answer('Qabul qilinmadi,bazadan o\'chiring')
    await state.finish()


@dp.message_handler(text='/delete_teacher', user_id=ADMINS[1])
@dp.message_handler(text='/delete_teacher', user_id=ADMINS[0])
async def delete_teacher_def(message: Message, state: FSMContext):
    await message.answer('Qaysi Fan ustozini o\'chramiz?', reply_markup=how_subject)
    await state.set_state('delete_teacher')


@dp.callback_query_handler(state='delete_teacher')
async def delete_teacher_defs(call: CallbackQuery, state: FSMContext):
    await call.message.delete()
    if call.data == 'it':
        db.delete_teacher(science=call.data)
        await call.answer("IT ustozi o'chirildi!", show_alert=True)
        await state.finish()
    elif call.data == 'ingliz':
        db.delete_teacher(science='english')
        await call.answer("english ustozi o'chirildi!", show_alert=True)
        await state.finish()

    elif call.data == 'math':
        db.delete_teacher(science=call.data)
        await call.answer("Math ustozi o'chirildi!", show_alert=True)
        await state.finish()
    await call.message.answer('Bosh Menu', reply_markup=bosh_menu)
    await state.finish()


@dp.message_handler(commands=['131_users'], user_id=ADMINS[0])
@dp.message_handler(commands=['131_users'], user_id=ADMINS[1])
async def send_file_131(message: Message, state: FSMContext):
    document = Document()
    document.add_heading('131-Maktab o\'quvchilari malumotlari', 0)
    p = document.add_paragraph()
    p.add_run('Ma\'lumotlar maxfiy saqlanadi!').bold = True
    users = db.select_user(maktab_raqami=131)
    for user in users:
        document.add_paragraph(f'Ism: {user[1]} ', style='List Number')
        document.add_paragraph(f'Telefon raqam: {user[2]}', style='List Number')
        document.add_paragraph(f'Maktabi soni: {user[3]}\n\n', style='List Number')
    document.add_page_break()
    document.save(f'{message.from_user.id}.docx')

    doc = InputFile(f'{message.from_user.id}.docx', filename=f'131-MAKTAB.docx')
    await ChatActions.upload_document(sleep=2)
    await message.answer_document(document=doc, caption='131-maktab,BAZA')
    os.remove(path=f"{message.from_user.id}.docx")


@dp.message_handler(commands=['165_users'], user_id=ADMINS[0])
@dp.message_handler(commands=['165_users'], user_id=ADMINS[1])
async def send_file_165(message: Message, state: FSMContext):
    document = Document()
    document.add_heading('165-Maktab o\'quvchilari malumotlari', 0)
    p = document.add_paragraph()
    p.add_run('Ma\'lumotlar maxfiy saqlanadi!').bold = True
    users = db.select_user(maktab_raqami=165)
    for user in users:
        document.add_paragraph(f'Ism: {user[1]} ', style='List Number')
        document.add_paragraph(f'Telefon raqam: {user[2]}', style='List Number')
        document.add_paragraph(f'Maktai soni: {user[3]}\n\n', style='List Number')
    document.add_page_break()
    document.save(f'{message.from_user.id}.docx')

    doc = InputFile(f'{message.from_user.id}.docx', filename=f'165-MAKTAB.docx')
    await ChatActions.upload_document(sleep=2)
    await message.answer_document(document=doc, caption='165-maktab,BAZA')
    os.remove(path=f"{message.from_user.id}.docx")


@dp.message_handler(commands=['330_users'], user_id=ADMINS[0])
@dp.message_handler(commands=['330_users'], user_id=ADMINS[1])
async def send_file_330(message: Message, state: FSMContext):
    document = Document()
    document.add_heading('330-Maktab o\'quvchilari malumotlari', 0)
    p = document.add_paragraph()
    p.add_run('Ma\'lumotlar maxfiy saqlanadi!').bold = True
    users = db.select_user(maktab_raqami=330)
    for user in users:
        document.add_paragraph(f'Ism: {user[1]} ', style='List Number')
        document.add_paragraph(f'Telefon raqam: {user[2]}', style='List Number')
        document.add_paragraph(f'Maktab soni: {user[3]}\n\n', style='List Number')
    document.add_page_break()
    document.save(f'{message.from_user.id}.docx')

    doc = InputFile(f'{message.from_user.id}.docx', filename=f'330-MAKTAB.docx')
    await ChatActions.upload_document(sleep=2)
    await message.answer_document(document=doc, caption='330-maktab,BAZASI')
    os.remove(path=f"{message.from_user.id}.docx")
